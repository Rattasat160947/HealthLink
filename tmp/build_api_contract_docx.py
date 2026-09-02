from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent / "CareKeeper_API_Contract.docx"

FONT_LATIN = "Calibri"
FONT_THAI = "Tahoma"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CODE_GRAY = "F7F7F7"
BORDER_GRAY = "B7B7B7"
MUTED = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, name: str = FONT_LATIN) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_THAI)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text: str = "", *, bold=False, size=11, color=None,
                  before=0, after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep_with_next
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_code_block(doc, text: str) -> None:
    for index, line in enumerate(text.splitlines()):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(4 if index == 0 else 0)
        p.paragraph_format.space_after = Pt(4 if index == len(text.splitlines()) - 1 else 0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_with_next = index < len(text.splitlines()) - 1
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CODE_GRAY)
        p_pr.append(shd)
        run = p.add_run(line if line else " ")
        set_run_font(run, size=9, name="Consolas")


def style_table(table, widths_dxa: list[int]) -> None:
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.10
                if col_index in (1, 2):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=(row_index == 0))


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    style_table(table, widths_dxa)
    add_paragraph(doc, after=2)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_THAI)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_THAI)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    set_run_font(run, size=9, color="777777")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def build() -> None:
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header_p = section.header.paragraphs[0]
    header_p.text = "CareKeeper | Backend API Contract"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    for run in header_p.runs:
        set_run_font(run, size=8.5, color="777777")
    add_page_number(section.footer.paragraphs[0])

    title = add_paragraph(doc, "ข้อเสนอรูปแบบ API ระบบ CareKeeper", bold=True,
                          size=22, after=4, keep_with_next=True)
    title.paragraph_format.space_before = Pt(4)
    add_paragraph(doc, "เอกสารสำหรับตรวจสอบและปรับ Backend API",
                  size=12, color="595959", after=12, keep_with_next=True)
    add_paragraph(doc, "สถานะ: ฉบับเสนอเพื่อพิจารณา | วันที่: 2 กันยายน 2026",
                  size=10, color="595959", after=16)

    add_paragraph(
        doc,
        "เอกสารนี้เสนอรูปแบบการรับส่งข้อมูลระหว่างเครื่อง CareKeeper และ Backend "
        "จำนวน 2 API ได้แก่ POST สำหรับบันทึกผลตรวจ และ GET สำหรับดึงประวัติผลตรวจย้อนหลัง",
        after=8,
    )

    add_heading(doc, "ข้อกำหนดร่วม", 1)
    add_table(
        doc,
        ["รายการ", "ค่าที่เสนอ"],
        [
            ["Authentication", "ส่ง API key ผ่าน Header ชื่อ api-key"],
            ["Content type", "application/json"],
            ["patient_id", "ตัวเลข 13 หลัก ไม่มีเครื่องหมายขีด"],
            ["measured_at", "รูปแบบ YYYY-MM-DD HH:MM:SS"],
            ["เขตเวลา", "Asia/Bangkok"],
        ],
        [2700, 6660],
    )

    add_heading(doc, "1. POST บันทึกผลตรวจสุขภาพ", 1)
    add_heading(doc, "Endpoint และ Headers", 2)
    add_code_block(doc, "POST /api/v2/device/add_health\nContent-Type: application/json\napi-key: <API_KEY>")

    add_heading(doc, "Request Body", 2)
    add_code_block(
        doc,
        '{\n'
        '  "mac": "1c:ce:51:9a:34:77",\n'
        '  "patient_id": "1234567890123",\n'
        '  "measured_at": "2026-09-01 20:15:33",\n'
        '  "sys": 120,\n'
        '  "dia": 80,\n'
        '  "pulse": 70,\n'
        '  "spo2": 98,\n'
        '  "temperature": 36.5\n'
        '}',
    )

    add_heading(doc, "รายละเอียด Request Body", 2)
    add_table(
        doc,
        ["Key", "ชนิด", "Required", "รายละเอียด"],
        [
            ["mac", "string", "ใช่", "MAC Address ของเครื่อง CareKeeper"],
            ["patient_id", "string", "ใช่", "เลขบัตรประชาชน 13 หลัก"],
            ["measured_at", "string", "ใช่", "วันและเวลาที่วัดเสร็จ"],
            ["sys", "integer", "ใช่", "ความดันตัวบน หน่วย mmHg"],
            ["dia", "integer", "ใช่", "ความดันตัวล่าง หน่วย mmHg"],
            ["pulse", "integer", "ใช่", "ชีพจร หน่วยครั้งต่อนาที"],
            ["spo2", "integer", "ใช่", "ออกซิเจนในเลือด หน่วยเปอร์เซ็นต์"],
            ["temperature", "number", "ใช่", "อุณหภูมิ หน่วยองศาเซลเซียส รองรับทศนิยม"],
        ],
        [1800, 1260, 1260, 5040],
    )

    add_heading(doc, "Response ที่เสนอ", 2)
    add_paragraph(doc, "เมื่อบันทึกสำเร็จ Backend ตอบ HTTP 200 หรือ 201 ตัวอย่าง:", after=4)
    add_code_block(doc, '{\n  "success": true,\n  "message": "บันทึกผลตรวจสุขภาพสำเร็จ"\n}')
    add_paragraph(
        doc,
        "โปรแกรม CareKeeper ถือว่า HTTP Status 200-299 สำเร็จ และไม่ผูกกับรูปแบบ Response Body",
        after=6,
    )

    add_heading(doc, "HTTP Status ที่เสนอ", 2)
    add_table(
        doc,
        ["Status", "ความหมาย"],
        [
            ["200 / 201", "บันทึกข้อมูลสำเร็จ"],
            ["400", "ข้อมูลไม่ครบหรือรูปแบบไม่ถูกต้อง"],
            ["401", "API key ไม่ถูกต้อง"],
            ["404", "ไม่พบผู้รับบริการหรืออุปกรณ์"],
            ["500", "Backend เกิดข้อผิดพลาด"],
        ],
        [1800, 7560],
    )

    add_heading(doc, "2. GET ดึงประวัติผลตรวจสุขภาพ", 1)
    add_heading(doc, "Endpoint และ Headers", 2)
    add_code_block(
        doc,
        "GET /api/v2/device/health_history?patient_id=1234567890123&mac=1c:ce:51:9a:34:77&limit=4\n"
        "Content-Type: application/json\n"
        "api-key: <API_KEY>",
    )
    add_paragraph(doc, "GET Request ไม่มี JSON Body โดยส่งค่าผ่าน Query Parameters", after=6)

    add_heading(doc, "Query Parameters", 2)
    add_table(
        doc,
        ["Key", "ชนิด", "Required", "รายละเอียด"],
        [
            ["patient_id", "string", "ใช่", "เลขบัตรประชาชน 13 หลัก"],
            ["mac", "string", "ใช่", "MAC Address ของเครื่อง CareKeeper"],
            ["limit", "integer", "ใช่", "จำนวนประวัติล่าสุด ปัจจุบันส่งค่า 4"],
        ],
        [1800, 1260, 1260, 5040],
    )

    add_heading(doc, "Response เมื่อสำเร็จ", 2)
    add_code_block(
        doc,
        '{\n'
        '  "success": true,\n'
        '  "data": [\n'
        '    {\n'
        '      "measured_at": "2026-09-01 20:15:33",\n'
        '      "sys": 120,\n'
        '      "dia": 80,\n'
        '      "pulse": 70,\n'
        '      "spo2": 98,\n'
        '      "temperature": 36.5\n'
        '    }\n'
        '  ]\n'
        '}',
    )
    add_paragraph(
        doc,
        "Backend ควรเรียง data จากรายการล่าสุดไปหาเก่าสุด และส่งกลับไม่เกินจำนวนที่ระบุใน limit",
        after=6,
    )

    add_heading(doc, "Response เมื่อไม่พบประวัติ", 2)
    add_code_block(doc, '{\n  "success": true,\n  "data": []\n}')

    add_heading(doc, "Key ภายใน data", 2)
    add_table(
        doc,
        ["Key", "ชนิด", "รายละเอียด"],
        [
            ["measured_at", "string", "วันและเวลาที่วัด"],
            ["sys", "integer", "ความดันตัวบน"],
            ["dia", "integer", "ความดันตัวล่าง"],
            ["pulse", "integer", "ชีพจร"],
            ["spo2", "integer", "ออกซิเจนในเลือด"],
            ["temperature", "number", "อุณหภูมิ รองรับทศนิยม"],
        ],
        [2340, 1800, 5220],
    )

    add_heading(doc, "3. ประเด็นที่ขอให้อาจารย์พิจารณา", 1)
    add_table(
        doc,
        ["ลำดับ", "ประเด็นที่ต้องยืนยัน"],
        [
            ["1", "ชื่อและ path ของ POST และ GET endpoint"],
            ["2", "ชื่อ Header สำหรับ API key ว่าจะใช้ api-key ตามที่เสนอหรือไม่"],
            ["3", "GET จำเป็นต้องใช้ทั้ง patient_id และ mac หรือใช้ patient_id อย่างเดียว"],
            ["4", "รูปแบบ Response Body และข้อความ error ที่ Backend จะกำหนด"],
            ["5", "กรณีเซนเซอร์ไม่มีค่า จะอนุญาตให้ส่ง null หรือให้ปฏิเสธ Request"],
            ["6", "ยืนยันเขตเวลาและรูปแบบ measured_at"],
        ],
        [1080, 8280],
    )

    add_paragraph(
        doc,
        "หมายเหตุ: โครงสร้างนี้เป็นข้อเสนอจากฝั่งโปรแกรม CareKeeper และสามารถปรับให้ตรงกับ Backend ที่อาจารย์กำหนดได้",
        bold=True,
        color=DARK_BLUE,
        before=8,
        after=0,
    )

    doc.core_properties.title = "ข้อเสนอรูปแบบ API ระบบ CareKeeper"
    doc.core_properties.subject = "POST และ GET Backend API Contract"
    doc.core_properties.author = "CareKeeper Project"
    doc.core_properties.keywords = "CareKeeper, API, Backend, POST, GET"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
